"""
Gera o arquivo Word do Fluxo de Estudo — Avaliação de Empresas (Mestrado Profissional)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Pinta o fundo de uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc, text: str, level: int, color: RGBColor = None, bold: bool = True):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.bold = bold
        if color:
            run.font.color.rgb = color
    return p


def add_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                   color: RGBColor = None, size: int = 11, indent: float = 0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def add_bullet(doc, text: str, level: int = 0, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    return p


def add_formula_box(doc, formula_lines: list, title: str = ""):
    """Cria um parágrafo de fórmula em fundo cinza claro."""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(f"  {title}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    for line in formula_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(f"  {line}")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)


def add_box_table(doc, header: str, header_color: str, content_items: list):
    """Cria uma tabela de 1 coluna com cabeçalho colorido e itens."""
    tbl = doc.add_table(rows=1 + len(content_items), cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(5.8)

    # Header
    hdr_cell = tbl.rows[0].cells[0]
    set_cell_bg(hdr_cell, header_color)
    p = hdr_cell.paragraphs[0]
    run = p.add_run(header)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Content
    for i, item in enumerate(content_items):
        cell = tbl.rows[i + 1].cells[0]
        p = cell.paragraphs[0]
        if isinstance(item, tuple):
            r1 = p.add_run(item[0])
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(item[1])
            r2.font.size = Pt(10)
        else:
            run = p.add_run(f"  {item}")
            run.font.size = Pt(10)
    doc.add_paragraph()


def add_two_col_table(doc, rows: list, header1: str = "", header2: str = ""):
    """Tabela de 2 colunas."""
    n = len(rows) + (1 if header1 else 0)
    tbl = doc.add_table(rows=n, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    start = 0
    if header1:
        c1, c2 = tbl.rows[0].cells
        set_cell_bg(c1, "1A5276")
        set_cell_bg(c2, "1A5276")
        for cell, text in [(c1, header1), (c2, header2)]:
            r = cell.paragraphs[0].add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
        start = 1

    for i, (col1, col2) in enumerate(rows):
        row = tbl.rows[start + i]
        if i % 2 == 0:
            set_cell_bg(row.cells[0], "EBF5FB")
            set_cell_bg(row.cells[1], "EBF5FB")
        for cell, text in [(row.cells[0], col1), (row.cells[1], col2)]:
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(10)
    doc.add_paragraph()


def divider(doc):
    doc.add_paragraph("─" * 70)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# ── CAPA ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FLUXO DE ESTUDO")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Avaliação de Empresas")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x21, 0x8A, 0xBE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mestrado Profissional em Economia")
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prof. Sérgio Jurandyr Machado")
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph()

# Índice resumido
p = doc.add_paragraph()
run = p.add_run("ESTRUTURA DO FLUXO DE ESTUDO")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

modulos = [
    ("MÓDULO 1", "Fundamentos do Valuation", "Aula 1"),
    ("MÓDULO 2", "Fluxo de Caixa Livre para a Firma (FCFF)", "Aulas 2–4"),
    ("MÓDULO 3", "Taxa de Desconto — WACC", "Aulas 2–4"),
    ("MÓDULO 4", "Ajustes Contábeis Relevantes para o Valuation", "Aulas 5–6"),
    ("MÓDULO 5", "Opções Reais e Valor da Flexibilidade", "Aulas 7–8"),
    ("MÓDULO 6", "Criação de Valor & Avaliação Relativa", "Transversal"),
    ("MÓDULO 7", "Revisão Geral e Exercícios Integrados", "Antes da Prova"),
]

for mod, titulo, aulas in modulos:
    p = doc.add_paragraph()
    r1 = p.add_run(f"  {mod}  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.size = Pt(10)
    # Simula destaque com tabs
    r2 = p.add_run(f"{titulo}  ")
    r2.font.size = Pt(11)
    r3 = p.add_run(f"[{aulas}]")
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 1 — FUNDAMENTOS DO VALUATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 1 — FUNDAMENTOS DO VALUATION", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Aulas de referência: Slide 02 (Aula 1)", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()

add_heading(doc, "1.1 O que é Valuation?", 2, RGBColor(0x21, 0x8A, 0xBE))
add_paragraph(doc,
    "Valuation é o processo de conversão de uma projeção em uma estimativa "
    "do valor de uma empresa ou de um ativo. Dois grandes métodos:")
add_bullet(doc, "Fluxo de Caixa Descontado (DCF / FCD): baseia-se na capacidade "
           "de geração de caixa futuro da empresa.", bold_prefix="")
add_bullet(doc, "Avaliação Relativa (Múltiplos): compara a empresa a pares de "
           "mercado por indicadores como EV/EBITDA, P/L etc.")
doc.add_paragraph()

add_heading(doc, "1.2 Tese de Investimento", 2, RGBColor(0x21, 0x8A, 0xBE))
add_paragraph(doc,
    '"O preço de um ativo converge para o seu valor justo."',
    italic=True, color=RGBColor(0x1A, 0x53, 0x76))
add_two_col_table(doc, [
    ("Fundo Long",        "Compra ativo subavaliado (preço < valor justo)"),
    ("Fundo Short",       "Vende ativo superavaliado (preço > valor justo)"),
    ("Fundo Long & Short","Opera nos dois sentidos simultaneamente"),
], "Estratégia", "Descrição")

add_heading(doc, "1.3 Eficiência de Mercado", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Hipótese dos Mercados Eficientes (HME): na forma semi-forte, "
           "toda informação pública é incorporada instantaneamente ao preço.")
add_bullet(doc, "Adaptive Markets Hypothesis (AMH): o preço converge para o "
           "valor justo, mas a velocidade pode variar por fricções de mercado "
           "(custos de transação, restrições à venda a descoberto) ou "
           "comportamento irracional (noise trader).")
doc.add_paragraph()

add_box_table(doc, "CONCEITOS-CHAVE DO MÓDULO 1", "1A5276", [
    ("Definição:", " Valuation = conversão de projeção em valor."),
    ("Premissa DCF:", " Valor = capacidade de geração de caixa futuro."),
    ("HME semi-forte:", " Informação pública → preço imediatamente."),
    ("Objetivo do analista:", " Gerar alfa: identificar o desvio entre preço e valor."),
    ("Citação de Damodaran:", ' "I would buy any company at the right price."'),
])

add_heading(doc, "Autoavaliação — Módulo 1", 2, RGBColor(0x7D, 0x6B, 0xE8))
perguntas_m1 = [
    "Qual a diferença entre preço e valor de um ativo?",
    "O que é HME em sua forma semi-forte? E a AMH?",
    "Por que 'gerar alfa é mais difícil do que ganhar as Olimpíadas'?",
    "Quais as três perguntas centrais do DCF?",
]
for q in perguntas_m1:
    add_bullet(doc, q)
    add_bullet(doc, "_" * 60, level=1)
    add_bullet(doc, "", level=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 2 — FCFF
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 2 — FLUXO DE CAIXA LIVRE PARA A FIRMA (FCFF)", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Aulas de referência: Aulas 2–4", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()

add_heading(doc, "2.1 Três Perguntas do DCF", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("Pergunta 1", "Qual o montante líquido de caixa gerado pelas atividades operacionais e como se distribui no tempo?"),
    ("Pergunta 2", "Qual deve ser a taxa de desconto?"),
    ("Pergunta 3", "Qual é o horizonte de análise (valor terminal)?"),
], "Nº", "Questão")

add_heading(doc, "2.2 Cálculo do FCFF Atual (Pergunta 1.a)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "EBIT",
    "(+)  Depreciação e Amortização",
    "( =) EBITDA  [Resolução CVM 156/2022]",
    "(+/-) Ajustes não recorrentes",
    "( =) EBITDA Ajustado",
    "( -) IR/CSLL estimado sobre o EBIT",
    "( -) CAPEX (investimento em capital fixo)",
    "( -) Variação da NCG (capital de giro não monetário)",
    "═════════════════════════════════════════",
    "( =) FCFF  em t",
], "Fórmula do FCFF")

doc.add_paragraph()

add_heading(doc, "2.3 CAPEX vs. OPEX", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("CAPEX", "Investimento que: (i) gera receita adicional; e/ou (ii) aumenta vida útil de ativo; e/ou (iii) expande capacidade. Concentrado no tempo. Parte custeada com capital de terceiros."),
    ("OPEX",  "Gasto do dia a dia (ex.: manutenção, salários operacionais). Já contemplado no EBIT. Mais frequente e recorrente."),
], "Tipo", "Característica")

add_heading(doc, "2.4 Necessidade de Capital de Giro (NCG)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "NCG = Estoques + Contas a Receber – Fornecedores",
    "",
    "Ciclo Financeiro = PME + PMR – PMP   (em dias de receita de vendas)",
    "  PME = Prazo Médio de Estoques     → quanto menor, melhor",
    "  PMR = Prazo Médio de Recebimento  → quanto menor, melhor",
    "  PMP = Prazo Médio de Pagamento    → quanto maior, melhor",
])
doc.add_paragraph()

add_heading(doc, "2.5 FCFF Atual como Base para Extrapolação (Pergunta 1.b)", 2,
            RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Verifique se há itens não recorrentes distorcendo o FCFF corrente.")
add_bullet(doc, "Avalie se o CAPEX do período é representativo ou pontual.")
add_bullet(doc, "Analise se a margem EBIT reflete condições normais de mercado.")
doc.add_paragraph()

add_heading(doc, "2.6 Taxa de Crescimento Esperada do FCFF (Pergunta 1.c)", 2,
            RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "g (crescimento) = ROIC × (CAPEX Líquido + ΔNCG) / NOPAT",
    "       ou",
    "g (crescimento) = Lucro Retido × ROE",
    "",
    "⚠  Atenção: g de longo prazo não deve superar o PIB Potencial ± ajustes de eficiência.",
])
doc.add_paragraph()

add_box_table(doc, "EXEMPLO RESOLVIDO — FCFF da Empresa XYZ (Polímeros)", "117A65", [
    ("Dados:", "EBIT = R$ 20 MM | Depreciação = R$ 1,5 MM | IR sobre EBIT = R$ 6,8 MM"),
    ("",       "Receita não recorrente = R$ 5,2 MM | ΔNCG = R$ 1 MM | CAPEX = R$ 2 MM"),
    ("FCFF 2025:", "20,0 + 1,5 – 5,2 – 6,8 – 1,0 – 2,0 = R$ 6,5 MM"),
    ("FCFF 2026:", "R$ 6,5 MM × 1,02 = R$ 6,63 MM (crescimento real de 2%)"),
])

add_heading(doc, "Autoavaliação — Módulo 2", 2, RGBColor(0x7D, 0x6B, 0xE8))
perguntas_m2 = [
    "Qual a diferença entre EBITDA e FCFF?",
    "Por que subtraímos CAPEX e ΔNCG no cálculo do FCFF?",
    "Por que adicionamos depreciação de volta ao EBIT?",
    "Como itens não recorrentes afetam a extrapolação do FCFF?",
    "Qual o risco de projetar g (crescimento) acima do PIB potencial de longo prazo?",
]
for q in perguntas_m2:
    add_bullet(doc, q)
    add_bullet(doc, "_" * 60, level=1)
    add_bullet(doc, "", level=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 3 — WACC
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 3 — TAXA DE DESCONTO: WACC", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Aulas de referência: Aulas 2–4", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()

add_heading(doc, "3.1 Custo de Capital Próprio (Ke) — CAPM", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "Ke  =  Rf  +  βi × (Rm – Rf)",
    "",
    "  Rf   = Taxa livre de risco",
    "  βi   = Beta da empresa i  (risco relativo ao mercado)",
    "  Rm–Rf = Prêmio de risco de mercado (ERP)",
])
doc.add_paragraph()

add_heading(doc, "3.2 Taxa Livre de Risco (Rf)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Requisitos: sem risco de default, sem risco de reinvestimento, "
           "duration compatível com o horizonte da análise.")
add_bullet(doc, "Sugestão do professor: NTN-B Principal com vencimento próximo a "
           "10 anos (taxa real). Já embute risco-país (Brasil).")
add_bullet(doc, "Em valuation real: Ke = Rf [NTN-B 10a] + βi × 5,5% a.a.")
doc.add_paragraph()

add_heading(doc, "3.3 Prêmio de Risco de Mercado (ERP)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("ERP Brasil", "Instável: varia com o horizonte; média geométrica do Ibovespa vs. SELIC sugere prêmio baixo em vários períodos."),
    ("ERP USA (S&P500)", "Média geométrica de Stocks − T.Bonds (1928–2024) ≈ 5,44%. Sugestão do professor: 5,5% a.a."),
    ("CAPM Global", "Ke = Rf[EUA] + Beta × (ERP[EUA] + Risco-País[BR]) + (INF_BR − INF_EUA). Avaliação em termos nominais."),
], "Abordagem", "Detalhamento")

add_heading(doc, "3.4 Beta (β)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "β  =  Cov(Ri, Rm) / σ²m",
    "",
    "  β < 1 → ativo menos volátil que o mercado",
    "  β > 1 → ativo mais volátil que o mercado",
    "",
    "Bottom-Up Beta: β setorial desalavancado → realavancado pela estrutura da empresa.",
    "  β_alavancado = β_desalavancado × [1 + (1 − t) × (D/E)]",
])
doc.add_paragraph()

add_heading(doc, "3.5 Custo de Capital de Terceiros (Kd)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "Kd = Rf + SOT   (spread over treasury)",
    "",
    "Custo Líquido de IR:  Kd × (1 − alíquota IR/CSLL)",
    "  → Juros são dedutíveis da base de cálculo do IR, gerando benefício fiscal.",
])
doc.add_paragraph()

add_heading(doc, "3.6 WACC", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "WACC = [E/(D+E)] × Ke  +  [D/(D+E)] × Kd × (1 − t)",
    "",
    "  E = Valor de mercado do capital próprio",
    "  D = Valor de mercado da dívida",
    "  t = Alíquota marginal de IR/CSLL",
    "",
    "Use pesos a valor de mercado (não contábeis).",
])
doc.add_paragraph()

add_box_table(doc, "EXEMPLO RESOLVIDO — Ke para Orizon (ORVR3) e Azul (AZUL4)", "117A65", [
    ("Data:", "22/10/2025 | NTN-B 2035 = 7,66% a.a. | ERP = 5,5% a.a."),
    ("Orizon:", "β = 0,21  →  Ke = 7,66 + 0,21 × 5,5 = 8,82% a.a."),
    ("Azul:",   "β = 1,73  →  Ke = 7,66 + 1,73 × 5,5 = 17,18% a.a."),
    ("Interpretação:", "Azul tem maior risco sistemático (setor aéreo vs. gestão de resíduos), exigindo maior retorno."),
])

add_heading(doc, "Autoavaliação — Módulo 3", 2, RGBColor(0x7D, 0x6B, 0xE8))
perguntas_m3 = [
    "Por que o Ke não é diretamente observável?",
    "Por que o Kd é mais barato que o Ke (em geral)?",
    "O que acontece com o WACC se a empresa aumenta sua alavancagem financeira?",
    "Por que usar a NTN-B e não a SELIC como taxa livre de risco em termos reais?",
    "O que é um 'bottom-up beta' e quando utilizá-lo?",
]
for q in perguntas_m3:
    add_bullet(doc, q)
    add_bullet(doc, "_" * 60, level=1)
    add_bullet(doc, "", level=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 4 — AJUSTES CONTÁBEIS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 4 — AJUSTES CONTÁBEIS RELEVANTES PARA O VALUATION", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Aulas de referência: Aulas 5–6", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()
add_paragraph(doc,
    "Estes itens são os 'Usual Suspects' — afetam (ou parecem afetar) o FCFF, "
    "mas sua influência real exige análise cuidadosa.",
    italic=True)
doc.add_paragraph()

# ── 4.1 Impairment ──
add_heading(doc, "4.1 Impairment (Perda por Desvalorização)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Conceito: quando Valor Contábil > Valor Recuperável, reconhece-se "
           "perda no resultado.")
add_bullet(doc, "Valor Recuperável = max(Valor Justo Líquido, Valor em Uso).")
add_bullet(doc, "Efeito sobre o FCFF: em si, NÃO afeta o caixa (despesa sem efeito "
           "caixa). Deve ser adicionado de volta, como a depreciação.")
add_bullet(doc, "Mas atenção: a reação do mercado ao impairment costuma ser negativa "
           "— o mercado entende que os fluxos futuros serão menores.")
add_bullet(doc, "Pergunta-chave: seus inputs de valuation já refletiam a redução de "
           "benefício do ativo impaired?")
add_paragraph(doc, "Casos discutidos em aula:", bold=True)
add_bullet(doc, "Petrobras (2015–2017): R$ 44,6 bi em impairment (Comperj R$ 21,8 bi).", level=1)
add_bullet(doc, "Grupo Soma / Hering: impairment do ágio por lei tributária (Lei nº 14.789/23).", level=1)
doc.add_paragraph()

# ── 4.2 Provisões ──
add_heading(doc, "4.2 Provisões (Contingências Passivas)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("Provável + mensurável",   "Provisionar no balanço (best estimate)"),
    ("Provável + não mensurável","Divulgar em notas explicativas"),
    ("Possível",                "Divulgar em notas explicativas"),
    ("Remota",                  "Não divulgar"),
], "Classificação", "Tratamento Contábil")
add_bullet(doc, "Relevância para valuation: contingências podem representar saídas "
           "futuras de caixa não capturadas no FCFF projetado.")
add_bullet(doc, "Mecanismos em M&A: Escrow Account e Earn-Out.")
add_bullet(doc,
    "Earn-Out: pagamento futuro ao vendedor condicionado ao atingimento de metas "
    "de desempenho da empresa adquirida. Soluciona divergências de expectativas.",
    bold_prefix="")
doc.add_paragraph()

# ── 4.3 Ágio (Goodwill) ──
add_heading(doc, "4.3 Ágio por Expectativa de Rentabilidade Futura (Goodwill)", 2,
            RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Surge em combinações de negócios: diferença entre o valor pago e "
           "o valor justo dos ativos líquidos adquiridos.")
add_bullet(doc, "Representa a estimativa de sinergias (redução de custo, ganho de "
           "mercado, etc.) em t₀.")
add_bullet(doc, "Para valuation: analise se o ágio é razoável — pode precisar ser "
           "ajustado para cima ou para baixo dependendo de sua análise de sinergias.")
doc.add_paragraph()

# ── 4.4 ARO ──
add_heading(doc, "4.4 ARO — Obrigação para Desmobilização de Ativos", 2,
            RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Quando o custo de retirada de serviço do bem é significativo e "
           "obrigatório, contabiliza-se o VP desse custo como ativo e provisão "
           "como passivo.")
add_bullet(doc, "O custo capitalizado é depreciado ao longo da vida útil do ativo.")
add_bullet(doc, "Para valuation: verifique se a ARO está sub ou superestimada. "
           "Caso de estudo: balanço da Vale antes e depois de Brumadinho (2018–2019).")
doc.add_paragraph()

# ── 4.5 POC ──
add_heading(doc, "4.5 Método POC (Percentage of Completion)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Reconhecimento de receita ao longo do tempo, proporcional ao "
           "avanço físico da obra/serviço.")
add_bullet(doc, "Aplicável somente se há transferência progressiva de controle ao "
           "cliente (IFRS 15 / CPC 47).")
add_bullet(doc, "Setor imobiliário: IASB esclareceu que o POC em geral NÃO é "
           "aderente ao IFRS 15 para incorporadoras residenciais no Brasil — "
           "receita deve ser reconhecida na entrega das chaves.")
add_bullet(doc, "Para valuation: cuidado com o timing do reconhecimento de receita; "
           "pode haver descolamento entre receita contábil e caixa recebido.")
doc.add_paragraph()

# ── 4.6 Hedge Accounting ──
add_heading(doc, "4.6 Hedge Accounting", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Permite diferir no patrimônio as variações de valor de instrumentos "
           "de hedge, reconhecendo-as apenas quando o item objeto é reconhecido "
           "no resultado.")
add_bullet(doc, "Caso Petrobras (2013–2015): área técnica da CVM entendeu que o "
           "hedge accounting foi usado para diferir perdas cambiais, desvirtuando "
           "a essência econômica. O Colegiado aceitou o recurso da empresa.")
add_bullet(doc, "Para valuation: identifique se o hedge accounting está mascarando "
           "perdas ou ganhos relevantes na DRE.")
doc.add_paragraph()

add_box_table(doc, "RESUMO — IMPACTO DOS 'USUAL SUSPECTS' NO FCFF", "76448A", [
    ("Impairment:",      "Sem efeito caixa direto. Adicionar de volta. Mas revise suas premissas de fluxo futuro."),
    ("Provisões:",       "Sem efeito caixa direto quando constituídas. O caixa sai apenas no pagamento."),
    ("ARO:",             "Custo capitalizado → depreciação aumenta EBITDA; pagamento futuro reduz caixa."),
    ("POC:",             "Receita pode antecipar ou diferir caixa → reconcilie sempre com o FCFF."),
    ("Hedge Acc.:",      "Pode distorcer DRE; analise o caixa efetivo da operação objeto."),
])

add_heading(doc, "Autoavaliação — Módulo 4", 2, RGBColor(0x7D, 0x6B, 0xE8))
perguntas_m4 = [
    "Por que o impairment não afeta o FCFF diretamente, mas afeta a percepção de valor?",
    "Qual a diferença entre escrow account e earn-out?",
    "Como uma ARO subestimada pode inflar artificialmente o valuation de uma empresa?",
    "Em que condições o POC é permitido pelo IFRS 15?",
    "Como identificar se o hedge accounting está sendo usado para diferir perdas?",
]
for q in perguntas_m4:
    add_bullet(doc, q)
    add_bullet(doc, "_" * 60, level=1)
    add_bullet(doc, "", level=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 5 — OPÇÕES REAIS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 5 — OPÇÕES REAIS E VALOR DA FLEXIBILIDADE", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Aulas de referência: Aulas 7–8", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()

add_heading(doc, "5.1 Conceito e Motivação", 2, RGBColor(0x21, 0x8A, 0xBE))
add_paragraph(doc,
    "O VPL tradicional ignora o valor das decisões gerenciais futuras "
    "(expandir, adiar, abandonar, mudar escala). As opções reais capturam "
    "esse valor de flexibilidade.")
add_paragraph(doc,
    "VPL Expandido = VPL Tradicional + Valor das Opções Reais")
doc.add_paragraph()

add_two_col_table(doc, [
    ("Opção de Expansão",  "Direito de ampliar o projeto se condições forem favoráveis. Análoga a uma opção de COMPRA (call)."),
    ("Opção de Abandono",  "Direito de encerrar o projeto e recuperar valor residual. Análoga a uma opção de VENDA (put)."),
    ("Opção de Adiamento", "Direito de postergar o início do investimento até obter mais informação."),
    ("Opção de Contração", "Direito de reduzir escala se condições deteriorarem."),
], "Tipo de Opção Real", "Descrição")

add_heading(doc, "5.2 Árvore Binomial", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "Em cada período, o valor do ativo sobe (×u) ou cai (×d):",
    "  u = fator de subida  (ex.: calculado a partir da volatilidade σ)",
    "  d = fator de descida (d = 1/u → árvore recombinante)",
    "",
    "Probabilidade Neutra ao Risco (q):",
    "  q  =  (1 + r – d) / (u – d)",
    "  r  =  taxa livre de risco por período",
    "",
    "Valor do nó pai = [q × Valor_sobe + (1−q) × Valor_desce] / (1 + r)",
])
doc.add_paragraph()

add_heading(doc, "5.3 Por que Usar a Probabilidade Neutra ao Risco?", 2,
            RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "O risco da opção muda a cada nó (depende do tempo, do preço de "
           "exercício e da volatilidade). Usar o WACC como taxa de desconto não é "
           "adequado para opções.")
add_bullet(doc, "A mudança da medida de probabilidade cria um mundo hipotético "
           "de neutralidade ao risco onde a taxa de desconto única é a taxa "
           "livre de risco.")
add_bullet(doc, "Esse 'truque matemático' (Girsanov) mantém os mesmos preços de "
           "mercado, mas permite descontar à taxa livre de risco.")
doc.add_paragraph()

add_heading(doc, "5.4 Black-Scholes para Opções Reais", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "Opção de Expansão (call):  c = S₀·N(d₁) − K·e^(−rT)·N(d₂)",
    "Opção de Abandono (put):   p = K·e^(−rT)·[1−N(d₂)] − S₀·[1−N(d₁)]",
    "",
    "  d₁ = [ln(S₀/K) + (r + σ²/2)·T] / (σ√T)",
    "  d₂ = d₁ − σ√T",
    "",
    "  S₀ = VP do fluxo de caixa do projeto (ex-CAPEX inicial)",
    "  K  = preço de exercício (custo do investimento futuro)",
    "  σ  = volatilidade (use proxy setorial se necessário)",
    "  r  = taxa livre de risco (contínuo)",
    "  T  = tempo até o vencimento (em anos)",
])
doc.add_paragraph()

add_heading(doc, "5.5 O que Afeta o Valor de uma Opção Real?", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("Preço do ativo objeto (S₀) ↑", "Call ↑   |   Put ↓"),
    ("Preço de exercício (K) ↑",     "Call ↓   |   Put ↑"),
    ("Volatilidade (σ) ↑",           "Call ↑   |   Put ↑"),
    ("Tempo até vencimento (T) ↑",   "Call ↑   |   Put ↑"),
    ("Taxa de juros (r) ↑",          "Call ↑   |   Put ↓"),
], "Parâmetro", "Efeito sobre o Valor da Opção")

add_heading(doc, "5.6 Exemplos Setoriais", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Indústria Cinematográfica: o primeiro filme é uma opção real de "
           "expansão — sequência só produzida se o primeiro for bem recebido. "
           "Preferência por franquias com menor incerteza de demanda.")
add_bullet(doc, "Indústria Farmacêutica: cada fase de P&D é análoga a uma opção "
           "de compra onde o custo atual é o prêmio e o custo de fases seguintes "
           "é o preço de exercício.")
add_bullet(doc, "Construção: opção de abandono com valor residual de venda a "
           "concorrente.")
doc.add_paragraph()

add_box_table(doc, "EXEMPLO RESOLVIDO — Opção de Abandono (Planta de Ração)", "76448A", [
    ("Dados:", "u=1,5 | d=0,5 | rf=5% a.a. | CAPEX=$ 1.000 | Preço de Exercício Abandono=$ 700"),
    ("Prob. neutra ao risco:", "q = (1,05 − 0,5) / (1,5 − 0,5) = 0,55  (55%)"),
    ("Valor sem opção (VPL A):", "Descontado ao WACC=30%; VPL negativo."),
    ("Valor com opção (VPL B):", "Maior, pois no cenário adverso a empresa vende a planta por $ 700."),
    ("Valor da Opção:", "VPL_B − VPL_A  = valor econômico da flexibilidade de abandono."),
])

add_heading(doc, "Autoavaliação — Módulo 5", 2, RGBColor(0x7D, 0x6B, 0xE8))
perguntas_m5 = [
    "Por que o VPL tradicional subestima o valor de projetos com flexibilidade gerencial?",
    "Por que não podemos usar o WACC para descontar os payoffs de uma opção real?",
    "O que é a probabilidade neutra ao risco e por que ela simplifica o cálculo?",
    "Aumento da volatilidade aumenta ou reduz o valor de uma opção de abandono? Por quê?",
    "Explique a analogia entre P&D farmacêutico e uma sequência de opções de compra.",
]
for q in perguntas_m5:
    add_bullet(doc, q)
    add_bullet(doc, "_" * 60, level=1)
    add_bullet(doc, "", level=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 6 — CRIAÇÃO DE VALOR & AVALIAÇÃO RELATIVA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 6 — CRIAÇÃO DE VALOR & AVALIAÇÃO RELATIVA", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Aulas de referência: transversal", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()

add_heading(doc, "6.1 Fontes de Criação de Valor", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("Ganhos de eficiência",          "Melhoria de margem, giro de ativos ou alavancagem → ROE maior."),
    ("Novos projetos com VPL > 0",    "TIR do projeto superior ao WACC: empresa cria riqueza para o acionista."),
    ("Alteração na estrutura de capital", "Uso estratégico de dívida (benefício fiscal) sem comprometer a saúde financeira."),
    ("Combinações de negócios (M&A)", "Sinergias operacionais e tributárias podem gerar VPL positivo para o adquirente."),
], "Alavanca", "Mecanismo")

add_heading(doc, "6.2 Decomposição do ROE (DuPont)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_formula_box(doc, [
    "ROE = Margem Líquida × Giro do Ativo × Multiplicador do PL",
    "",
    "  Margem Líquida = Lucro Líquido / Receita",
    "  Giro do Ativo  = Receita / Ativo Total",
    "  Multiplicador  = Ativo Total / PL   (alavancagem)",
])
doc.add_paragraph()

add_heading(doc, "6.3 Caso Prático: Mobly & Tok&Stok (2024)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Sinergias operacionais: tecnologia/logística da Mobly + expertise de "
           "produto da Tok&Stok → Bain&Co estimou +R$ 80–135 MM de caixa/ano.")
add_bullet(doc, "Estrutura de capital: Mobly assume dívida de R$ 450 MM da Tok&Stok, "
           "avaliada por apenas 10% do seu pico de valor.")
add_bullet(doc, "Earn-out embutido na operação: troca de ações com lock-up de 24 meses.")
doc.add_paragraph()

add_heading(doc, "6.4 Avaliação Relativa (Múltiplos)", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("EV/EBITDA",  "Enterprise Value / EBITDA. Múltiplo de operações, neutro à estrutura de capital."),
    ("P/L (P/E)",  "Preço / Lucro por Ação. Reflete expectativas de crescimento."),
    ("P/VPA",      "Preço / Valor Patrimonial por Ação. Relevante para bancos e financeiras."),
    ("EV/Receita", "Usado em empresas pré-lucro ou de alto crescimento."),
], "Múltiplo", "Aplicação")
add_bullet(doc, "Premissa central: 'ativos semelhantes devem valer o mesmo.'")
add_bullet(doc, "Risco: múltiplos de mercado embitem otimismo ou pessimismo sistêmico "
           "do momento.")
doc.add_paragraph()

add_heading(doc, "Autoavaliação — Módulo 6", 2, RGBColor(0x7D, 0x6B, 0xE8))
perguntas_m6 = [
    "Em que condições um novo projeto cria valor para o acionista?",
    "Como o uso de dívida pode criar valor (benefício fiscal) e, ao mesmo tempo, destruí-lo (custos de dificuldade financeira)?",
    "Quais são os riscos de usar apenas múltiplos para avaliar uma empresa?",
    "Como as sinergias de uma fusão devem ser refletidas no valuation do alvo?",
]
for q in perguntas_m6:
    add_bullet(doc, q)
    add_bullet(doc, "_" * 60, level=1)
    add_bullet(doc, "", level=1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 7 — REVISÃO GERAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "MÓDULO 7 — REVISÃO GERAL E EXERCÍCIOS INTEGRADOS", 1,
            RGBColor(0x1A, 0x53, 0x76))
add_paragraph(doc, "Use este módulo na semana anterior à avaliação.", italic=True,
              color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()

add_heading(doc, "7.1 Checklist de Valuation (DCF Completo)", 2, RGBColor(0x21, 0x8A, 0xBE))
checklist = [
    "[ ] Calculei o FCFF atual e removi itens não recorrentes?",
    "[ ] Meu FCFF é uma base razoável para extrapolação?",
    "[ ] A taxa de crescimento de longo prazo é compatível com o PIB potencial?",
    "[ ] Calculei o Ke pelo CAPM usando NTN-B 10a e ERP de 5,5%?",
    "[ ] Usei o beta adequado (histórico ou bottom-up)?",
    "[ ] Calculei o Kd líquido de IR?",
    "[ ] Usei pesos do WACC a valor de mercado?",
    "[ ] Identifiquei e ajustei impairments, provisões e ARO relevantes?",
    "[ ] Verifiquei o POC / hedge accounting nos ajustes de receita?",
    "[ ] Considerei opções reais relevantes (expansão, abandono, adiamento)?",
    "[ ] Realizei análise de sensibilidade (g e WACC)?",
    "[ ] Comparei o resultado com múltiplos de mercado (sanity check)?",
]
for item in checklist:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    run.font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, "7.2 Mapa de Conexões", 2, RGBColor(0x21, 0x8A, 0xBE))
add_two_col_table(doc, [
    ("FCFF ↑",                "VPL do DCF ↑   (ceteris paribus)"),
    ("WACC ↑",                "VPL do DCF ↓   (ceteris paribus)"),
    ("g de longo prazo ↑",    "Valor Terminal ↑ → maior sensibilidade"),
    ("β ↑",                   "Ke ↑ → WACC ↑ → VPL ↓"),
    ("Alavancagem ↑",         "Kd(1−t) mais barato → WACC pode ↓; mas risco financeiro ↑"),
    ("Volatilidade ↑",        "Opções Reais ↑ (call e put)"),
    ("Impairment",            "Sem efeito direto no FCFF; revisar premissas de crescimento"),
    ("POC antecipado",        "Receita contábil antes do caixa → FCFF não afetado"),
], "Variável / Evento", "Efeito no Valuation")

add_heading(doc, "7.3 Exercício Integrado", 2, RGBColor(0x21, 0x8A, 0xBE))
add_paragraph(doc,
    "A empresa ABC atua no setor de logística. Com base nos dados abaixo, "
    "estime o FCFF, calcule o Ke, o WACC e discuta se existe opção real relevante:")
add_two_col_table(doc, [
    ("EBIT",              "R$ 50 MM"),
    ("Depreciação",       "R$ 8 MM"),
    ("Receita não recorr.","R$ 3 MM"),
    ("IR sobre EBIT",     "R$ 15 MM"),
    ("CAPEX",             "R$ 12 MM"),
    ("ΔNCG",              "R$ 4 MM"),
    ("Beta (histórico)",  "0,85"),
    ("NTN-B 10a",         "7,5% a.a."),
    ("ERP",               "5,5% a.a."),
    ("Kd bruto",          "11% a.a."),
    ("Alíquota IR",       "34%"),
    ("E/(D+E)",           "60%"),
    ("g projetado",       "3% a.a. real"),
], "Dado", "Valor")
add_paragraph(doc, "Resolução:", bold=True)
for linha in [
    "FCFF = 50 + 8 − 3 − 15 − 12 − 4  =  R$ 24 MM",
    "Ke   = 7,5 + 0,85 × 5,5           = 12,175% a.a.",
    "Kd líquido = 11 × (1 − 0,34)      =  7,26% a.a.",
    "WACC = 0,60 × 12,175 + 0,40 × 7,26 = 10,21% a.a.",
    "Valor Terminal = FCFF × (1+g) / (WACC−g) = 24×1,03 / (0,1021−0,03) = ~R$ 343 MM",
    "Discuta: existe opção real de expansão (nova rota/terminal) relevante?",
]:
    add_bullet(doc, linha, level=1)
doc.add_paragraph()

add_heading(doc, "7.4 Dicas para a Avaliação (Estudos Dirigidos e Seminário)", 2,
            RGBColor(0x21, 0x8A, 0xBE))
add_bullet(doc, "Leia os relatórios de resultado (ITR/DFP) com atenção às notas "
           "explicativas sobre contingências, impairment e segmentos.")
add_bullet(doc, "Em estudos de caso, sempre justifique os ajustes ao FCFF.")
add_bullet(doc, "No seminário, contextualize o artigo e explicite a questão de pesquisa "
           "nos dois primeiros slides (conforme diretriz do professor).")
add_bullet(doc, "Tempo de apresentação: ~30 minutos; upload até quinta-feira anterior às 23h59.")

doc.add_paragraph()
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fluxo de Estudo gerado com base nos Slides das Aulas 1–8 — Avaliação de Empresas")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Salvar ────────────────────────────────────────────────────────────────────
output_path = "/home/user/estado-economia-dashboard/Fluxo_de_Estudo_Avaliacao_de_Empresas.docx"
doc.save(output_path)
print(f"Arquivo salvo: {output_path}")
